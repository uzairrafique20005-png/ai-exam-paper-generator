"""
ExamCraft AI – Exam Paper Generator
Powered by Groq API (llama-3.1-8b-instant)
Full redesign: light theme, animations, PDF export, proper Bloom's taxonomy formatting
"""

import os
import io
import re
import streamlit as st
from groq import Groq
import PyPDF2
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    Table, TableStyle, PageBreak
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

# ─────────────────────────────────────────────
# COLOUR PALETTE (shared across app + exports)
# ─────────────────────────────────────────────
TEAL        = "#1A7A6E"
TEAL_LIGHT  = "#E8F5F3"
TEAL_MID    = "#2AA090"
AMBER       = "#E07B39"
AMBER_LIGHT = "#FDF3EC"
SLATE       = "#2D3748"
SLATE_LIGHT = "#718096"
WHITE       = "#FFFFFF"
OFF_WHITE   = "#F7FAFA"
BORDER      = "#D1E8E4"

# ReportLab colour objects
RL_TEAL        = colors.HexColor(TEAL)
RL_TEAL_LIGHT  = colors.HexColor(TEAL_LIGHT)
RL_AMBER       = colors.HexColor(AMBER)
RL_AMBER_LIGHT = colors.HexColor(AMBER_LIGHT)
RL_SLATE       = colors.HexColor(SLATE)
RL_SLATE_LIGHT = colors.HexColor(SLATE_LIGHT)
RL_WHITE       = colors.white
RL_BORDER      = colors.HexColor(BORDER)

# python-docx RGBColor objects
D_TEAL  = RGBColor(0x1A, 0x7A, 0x6E)
D_AMBER = RGBColor(0xE0, 0x7B, 0x39)
D_SLATE = RGBColor(0x2D, 0x37, 0x48)
D_GRAY  = RGBColor(0x71, 0x80, 0x96)

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="ExamCraft AI",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
# CUSTOM CSS – light theme, animations, always readable
# ─────────────────────────────────────────────
st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');

/* ── Force light base ── */
html, body, [data-testid="stAppViewContainer"], .stApp {{
    background-color: #F0F4F8 !important;
    color: {SLATE} !important;
    font-family: 'Inter', sans-serif !important;
}}

/* ── Sidebar: deep dark premium ── */
section[data-testid="stSidebar"] {{
    background: linear-gradient(180deg, #0F2027 0%, #1A2F38 50%, #0F2027 100%) !important;
    border-right: 1px solid rgba(26,122,110,0.3) !important;
}}

/* White text for sidebar elements */
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3,
section[data-testid="stSidebar"] h4,
section[data-testid="stSidebar"] > div p,
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] .stSelectbox label,
section[data-testid="stSidebar"] .stNumberInput label,
section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] > p,
section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] > h2,
section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] > h3 {{
    color: {WHITE} !important;
    font-family: 'Inter', sans-serif !important;
}}

/* Radio button labels in sidebar */
section[data-testid="stSidebar"] .stRadio label p,
section[data-testid="stSidebar"] .stRadio label span:not([style]) {{
    color: {WHITE} !important;
}}

/* Caption / small text in sidebar */
section[data-testid="stSidebar"] .stCaption,
section[data-testid="stSidebar"] .stCaption p {{
    color: rgba(255,255,255,0.5) !important;
}}

/* Sidebar selectbox + number input */
section[data-testid="stSidebar"] [data-baseweb="select"] > div,
section[data-testid="stSidebar"] [data-baseweb="input"] > div {{
    background: rgba(255,255,255,0.08) !important;
    border-color: rgba(26,122,110,0.5) !important;
    color: {WHITE} !important;
    border-radius: 10px !important;
}}
section[data-testid="stSidebar"] [data-baseweb="select"] svg {{
    fill: rgba(255,255,255,0.6) !important;
}}

/* Number input text in sidebar */
section[data-testid="stSidebar"] input[type="number"] {{
    color: {WHITE} !important;
    background: transparent !important;
}}

/* ── Main content – always dark readable text ── */
.main .block-container p,
.main .block-container li,
.main .block-container label,
.stMarkdown p, .stMarkdown li {{
    color: {SLATE} !important;
    font-family: 'Inter', sans-serif !important;
}}
h1, h2, h3, h4 {{ color: {SLATE} !important; font-family: 'Inter', sans-serif !important; }}

/* ── Radio button labels – main area ── */
.stRadio label,
.stRadio label p,
.stRadio label span,
div[data-testid="stRadio"] label,
div[data-testid="stRadio"] label p {{
    color: {SLATE} !important;
    font-weight: 600 !important;
    font-family: 'Inter', sans-serif !important;
}}
div[data-testid="stRadio"] > label > div > p {{
    color: {SLATE} !important;
}}

/* ── Metric labels ── */
[data-testid="stMetricLabel"]  {{ color: {SLATE_LIGHT} !important; font-size: 0.78rem !important; letter-spacing: 0.5px; }}
[data-testid="stMetricValue"]  {{ color: {TEAL} !important; font-weight: 800 !important; font-size: 1.6rem !important; }}

/* ── Hero banner: sleek premium ── */
.hero-banner {{
    background: linear-gradient(135deg, #0F2027 0%, #1A4A42 50%, {TEAL} 100%);
    border-radius: 20px;
    padding: 3rem 2.5rem 2.5rem 2.5rem;
    margin-bottom: 2rem;
    animation: fadeSlideDown 0.6s cubic-bezier(0.22,1,0.36,1) forwards;
    box-shadow: 0 20px 60px rgba(15,32,39,0.35), 0 0 0 1px rgba(26,122,110,0.2);
    position: relative;
    overflow: hidden;
}}
.hero-banner::before {{
    content: '';
    position: absolute;
    top: -50%;
    right: -10%;
    width: 400px;
    height: 400px;
    background: radial-gradient(circle, rgba(26,122,110,0.25) 0%, transparent 70%);
    border-radius: 50%;
}}
.hero-eyebrow {{
    font-size: 0.72rem;
    font-weight: 700;
    letter-spacing: 3px;
    text-transform: uppercase;
    color: {TEAL_MID} !important;
    margin: 0 0 0.6rem 0;
}}
.hero-title {{
    font-size: 2.8rem;
    font-weight: 900;
    color: {WHITE} !important;
    margin: 0 0 0.5rem 0;
    letter-spacing: -1.5px;
    line-height: 1.1;
}}
.hero-title span {{ color: {TEAL_MID} !important; }}
.hero-sub {{
    color: rgba(255,255,255,0.65) !important;
    font-size: 1rem;
    margin: 0;
    font-weight: 400;
    max-width: 560px;
}}

/* ── Section cards: glass morphism ── */
.section-card {{
    background: {WHITE};
    border: 1px solid rgba(209,232,228,0.8);
    border-radius: 16px;
    padding: 1.8rem 2rem;
    margin-bottom: 1.4rem;
    box-shadow: 0 4px 24px rgba(26,122,110,0.07), 0 1px 3px rgba(0,0,0,0.04);
    animation: fadeIn 0.5s cubic-bezier(0.22,1,0.36,1) forwards;
    position: relative;
}}
.section-card::before {{
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 3px;
    background: linear-gradient(90deg, {TEAL}, {TEAL_MID}, transparent);
    border-radius: 16px 16px 0 0;
}}
.section-label {{
    font-size: 0.68rem;
    font-weight: 800;
    letter-spacing: 2.5px;
    text-transform: uppercase;
    color: {TEAL} !important;
    margin-bottom: 0.3rem;
    display: flex;
    align-items: center;
    gap: 6px;
}}
.section-label::after {{
    content: '';
    flex: 1;
    height: 1px;
    background: linear-gradient(90deg, {BORDER}, transparent);
    margin-left: 8px;
}}
.section-title {{
    font-size: 1.2rem;
    font-weight: 700;
    color: {SLATE} !important;
    margin: 0 0 1.2rem 0;
}}

/* ── Question counter cards ── */
.q-counter-grid {{
    display: grid;
    grid-template-columns: repeat(3, 1fr);
    gap: 1rem;
    margin: 1rem 0 1.5rem 0;
}}
.q-counter-card {{
    background: linear-gradient(135deg, #F8FBFA, {WHITE});
    border: 1.5px solid {BORDER};
    border-radius: 12px;
    padding: 1rem 1.2rem;
    text-align: center;
    box-shadow: 0 2px 8px rgba(26,122,110,0.05);
}}
.q-counter-card.mcq {{ border-top: 3px solid {TEAL}; }}
.q-counter-card.short {{ border-top: 3px solid {AMBER}; }}
.q-counter-card.long {{ border-top: 3px solid #7E22CE; }}
.q-counter-label {{
    font-size: 0.7rem;
    font-weight: 700;
    letter-spacing: 1.5px;
    text-transform: uppercase;
    color: {SLATE_LIGHT};
    margin-bottom: 0.4rem;
}}

/* ── File uploader – fix dark background text ── */
[data-testid="stFileUploader"] {{
    background: {WHITE} !important;
    border: 2px dashed {BORDER} !important;
    border-radius: 12px !important;
    padding: 0.5rem !important;
}}
[data-testid="stFileUploader"] label,
[data-testid="stFileUploader"] span,
[data-testid="stFileUploader"] p,
[data-testid="stFileUploaderDropzone"] span,
[data-testid="stFileUploaderDropzone"] p {{
    color: {SLATE} !important;
    font-weight: 600 !important;
}}
[data-testid="stFileUploaderDropzone"] {{
    background: {WHITE} !important;
    border-radius: 10px !important;
}}
[data-testid="stFileUploaderDropzone"] button {{
    background: {WHITE} !important;
    border: 1.5px solid {TEAL} !important;
    color: {TEAL} !important;
    border-radius: 8px !important;
    font-weight: 700 !important;
}}
[data-testid="stFileUploaderDropzone"] button:hover {{
    background: {TEAL} !important;
    color: {WHITE} !important;
}}
/* Small hint text in uploader */
[data-testid="stFileUploaderDropzone"] small,
[data-testid="stFileUploaderDropzone"] .st-emotion-cache-1b0udgb {{
    color: {SLATE_LIGHT} !important;
}}

/* ── Expander: fix white-on-white preview text ── */
details {{
    background: {WHITE} !important;
    border: 1px solid {BORDER} !important;
    border-radius: 12px !important;
    overflow: hidden !important;
}}
details summary {{
    background: {TEAL_LIGHT} !important;
    border-radius: 0 !important;
    padding: 0.85rem 1.2rem !important;
    font-weight: 700 !important;
    color: {TEAL} !important;
    cursor: pointer !important;
    font-size: 0.92rem !important;
    border-bottom: 1px solid {BORDER} !important;
}}
details[open] summary {{
    border-radius: 0 !important;
}}
/* ALL text inside expander content must be dark */
details > div,
details > div *,
[data-testid="stExpander"] > div,
[data-testid="stExpander"] > div * {{
    color: {SLATE} !important;
    background-color: transparent !important;
}}
/* Specifically fix st.text() inside expander */
[data-testid="stExpander"] pre,
[data-testid="stExpander"] code,
.stExpander pre,
.stExpander code {{
    color: {SLATE} !important;
    background: #F8FAFA !important;
    border: 1px solid {BORDER} !important;
    border-radius: 8px !important;
    padding: 1rem !important;
    font-size: 0.85rem !important;
    line-height: 1.7 !important;
    white-space: pre-wrap !important;
}}

/* ── Exam output card ── */
.exam-output {{
    background: {WHITE};
    border-left: 4px solid {TEAL};
    border-radius: 12px;
    padding: 2rem 2.2rem;
    margin: 1rem 0;
    box-shadow: 0 6px 30px rgba(26,122,110,0.08);
    line-height: 1.85;
    color: {SLATE} !important;
}}
.exam-output * {{ color: {SLATE} !important; }}

/* ── Primary generate button ── */
.stButton > button {{
    background: linear-gradient(135deg, #0F2027 0%, {TEAL} 100%) !important;
    color: {WHITE} !important;
    border: none !important;
    border-radius: 12px !important;
    padding: 0.75rem 2rem !important;
    font-size: 1rem !important;
    font-weight: 700 !important;
    box-shadow: 0 6px 20px rgba(15,32,39,0.3) !important;
    transition: all 0.2s cubic-bezier(0.22,1,0.36,1) !important;
    letter-spacing: 0.3px !important;
    font-family: 'Inter', sans-serif !important;
}}
.stButton > button:hover {{
    box-shadow: 0 10px 30px rgba(15,32,39,0.4) !important;
    transform: translateY(-3px) !important;
}}
.stButton > button:active {{ transform: translateY(0) !important; }}

/* ── Download buttons ── */
.stDownloadButton > button {{
    border-radius: 12px !important;
    font-weight: 700 !important;
    transition: all 0.2s ease !important;
    border: 2px solid {TEAL} !important;
    color: {TEAL} !important;
    background: {WHITE} !important;
    font-family: 'Inter', sans-serif !important;
}}
.stDownloadButton > button:hover {{
    background: {TEAL} !important;
    color: {WHITE} !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 20px rgba(26,122,110,0.3) !important;
}}

/* ── Radio ── */
.stRadio label {{ color: {SLATE} !important; font-weight: 600 !important; }}
.stRadio label p, .stRadio label span {{ color: {SLATE} !important; }}

/* ── Text areas ── */
.stTextArea textarea {{
    border-radius: 10px !important;
    border: 1.5px solid {BORDER} !important;
    background: {WHITE} !important;
    color: {SLATE} !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 0.95rem !important;
    line-height: 1.6 !important;
}}
.stTextArea textarea:focus {{
    border-color: {TEAL} !important;
    box-shadow: 0 0 0 3px rgba(26,122,110,0.12) !important;
}}

/* ── Number inputs in main ── */
.stNumberInput input {{
    border-radius: 10px !important;
    border: 1.5px solid {BORDER} !important;
    background: {WHITE} !important;
    color: {SLATE} !important;
    font-weight: 700 !important;
    font-size: 1.1rem !important;
    text-align: center !important;
}}
.stNumberInput input:focus {{
    border-color: {TEAL} !important;
    box-shadow: 0 0 0 3px rgba(26,122,110,0.12) !important;
}}

/* ── Alerts ── */
.stAlert {{ border-radius: 12px !important; }}

/* ── Success message ── */
.stSuccess {{ 
    background: linear-gradient(135deg, #F0FDF4, #E6FAF5) !important; 
    border: 1px solid #86EFAC !important; 
    border-radius: 12px !important; 
}}

/* ── Selectbox ── */
[data-baseweb="select"] > div {{
    border-radius: 10px !important;
    border: 1.5px solid {BORDER} !important;
    background: {WHITE} !important;
}}

/* ── Animations ── */
@keyframes fadeSlideDown {{
    from {{ opacity: 0; transform: translateY(-30px); }}
    to   {{ opacity: 1; transform: translateY(0); }}
}}
@keyframes fadeIn {{
    from {{ opacity: 0; transform: translateY(12px); }}
    to   {{ opacity: 1; transform: translateY(0); }}
}}
@keyframes pulse {{
    0%, 100% {{ opacity: 1; }}
    50% {{ opacity: 0.6; }}
}}

/* ── Hide Streamlit chrome ── */
#MainMenu, footer {{ visibility: hidden; }}

/* ── Block container padding ── */
.block-container {{ padding-top: 2rem !important; padding-bottom: 3rem !important; }}
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# API KEY LOADER
# ─────────────────────────────────────────────
def load_api_key() -> str:
    try:
        return st.secrets["GROQ_API_KEY"]
    except (KeyError, FileNotFoundError):
        key = os.environ.get("GROQ_API_KEY", "")
        if not key:
            st.error("⚠️ Groq API key not found. Add to `.streamlit/secrets.toml` or set GROQ_API_KEY env var.")
            st.stop()
        return key


# ─────────────────────────────────────────────
# FILE EXTRACTOR
# ─────────────────────────────────────────────
MAX_CONTEXT_CHARS = 3000

def extract_file_content(uploaded_file) -> str:
    try:
        if uploaded_file.type == "application/pdf":
            reader = PyPDF2.PdfReader(uploaded_file)
            pages = [page.extract_text() or "" for page in reader.pages]
            raw_text = "\n".join(pages).strip()
        else:
            raw_text = uploaded_file.read().decode("utf-8", errors="ignore").strip()
    except Exception as e:
        st.warning(f"File reading error: {e}.")
        return ""
    if len(raw_text) > MAX_CONTEXT_CHARS:
        st.warning("📄 Document is large. Only the first ~500 words were used to stay within API limits.")
        return raw_text[:MAX_CONTEXT_CHARS]
    return raw_text


# ─────────────────────────────────────────────
# PROMPT – enforces Bloom's tags + clean delimiters
# ─────────────────────────────────────────────
def build_prompt(topic_context, grade, difficulty, duration, total_marks, source_type, num_mcq=5, num_short=3, num_long=1):
    context_instruction = (
        f"Base the entire exam STRICTLY on this uploaded source material:\n\n{topic_context}"
        if source_type == "Upload Source Material"
        else f"Cover the following subject/topics:\n\n{topic_context}"
    )
    total_q = num_mcq + num_short + num_long

    # Build dynamic section instructions
    mcq_section = ""
    if num_mcq > 0:
        mcq_section = f"""
### SECTION A — MULTIPLE CHOICE QUESTIONS ({num_mcq} Questions)

""" + "\n\n".join([
            f"Q{i+1}. [BLOOM: Knowledge] (2 Marks)\nQuestion text here?\n   A) Option one\n   B) Option two\n   C) Option three\n   D) Option four"
            for i in range(min(num_mcq, 3))
        ]) + f"\n... (continue for all {num_mcq} MCQs)\n"

    short_start = num_mcq + 1
    short_section = ""
    if num_short > 0:
        short_section = f"""
### SECTION B — SHORT ANSWER QUESTIONS ({num_short} Questions)

""" + "\n\n".join([
            f"Q{short_start+i}. [BLOOM: Application] (5 Marks)\nQuestion text here?"
            for i in range(min(num_short, 3))
        ]) + f"\n... (continue for all {num_short} short questions)\n"

    long_start = num_mcq + num_short + 1
    long_section = ""
    if num_long > 0:
        long_section = f"""
### SECTION C — LONG / ESSAY QUESTIONS ({num_long} Questions)

""" + "\n\n".join([
            f"Q{long_start+i}. [BLOOM: Synthesis] (20 Marks)\nEssay question text here?"
            for i in range(min(num_long, 2))
        ]) + f"\n... (continue for all {num_long} long questions)\n"

    return f"""You are an expert curriculum designer. Produce a complete exam paper and answer key.

CONFIGURATION:
- Grade/Semester: {grade}
- Difficulty: {difficulty}
- Duration: {duration} minutes
- Total Marks: {total_marks}
- Total Questions: {total_q} ({num_mcq} MCQs, {num_short} short, {num_long} long/essay)
- {context_instruction}

OUTPUT FORMAT — follow this structure EXACTLY, generating EXACTLY the number of questions specified:

---EXAM START---

## EXAM PAPER
{mcq_section}{short_section}{long_section}
---ANSWER KEY---

## ANSWER KEY & GRADING RUBRIC

### MCQ Answers
(List correct answers and brief explanations for all {num_mcq} MCQs)

### Short Answer — Model Answers
(Model answers for all {num_short} short questions, 3-5 sentences each)

### Long Answer — Grading Rubric
| Criterion | Excellent (5) | Good (3-4) | Needs Work (1-2) |
|---|---|---|---|
| Argument / Thesis | Clear original thesis | Adequate thesis | Weak or missing |
| Evidence & Examples | 3+ relevant examples | 1-2 examples | Vague or missing |
| Structure & Clarity | Logical flow | Mostly clear | Disorganised |
| Critical Thinking | Deep insight | Some analysis | Surface level |

RULES:
- Generate EXACTLY {num_mcq} MCQs, {num_short} short questions, and {num_long} long/essay questions.
- EVERY question MUST have [BLOOM: Level] tag exactly as shown.
- Use ---EXAM START--- and ---ANSWER KEY--- as the only section separators.
- Distribute {total_marks} marks proportionally across all {total_q} questions.
- No extra commentary — output only the exam and answer key.
"""


# ─────────────────────────────────────────────
# GROQ API
# ─────────────────────────────────────────────
def generate_exam(prompt: str, api_key: str) -> str:
    client = Groq(api_key=api_key)
    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=4096,
    )
    return response.choices[0].message.content


# ─────────────────────────────────────────────
# SPLIT TEXT
# ─────────────────────────────────────────────
BLOOM_RE = re.compile(r'\[BLOOM:\s*([A-Za-z]+)\]')

def split_exam_and_key(full_text: str):
    marker = "---ANSWER KEY---"
    if marker in full_text:
        idx = full_text.find(marker)
        return full_text[:idx].strip(), full_text[idx:].strip()
    for kw in ["ANSWER KEY", "═══"]:
        if kw in full_text.upper():
            idx = full_text.upper().find(kw)
            ls  = full_text.rfind("\n", 0, idx)
            return full_text[:ls].strip(), full_text[ls:].strip()
    return full_text.strip(), ""


# ─────────────────────────────────────────────
# HTML RENDERER for on-screen display
# ─────────────────────────────────────────────
BLOOM_COLOURS = {
    "knowledge":     ("#EEF2FF", "#3730A3"),
    "comprehension": ("#F0FDF4", "#166534"),
    "application":   ("#FFF7ED", "#C2410C"),
    "analysis":      ("#FDF4FF", "#7E22CE"),
    "synthesis":     ("#FFF1F2", "#BE123C"),
    "evaluation":    ("#ECFEFF", "#0E7490"),
}

def bloom_badge_html(level: str) -> str:
    key = level.lower()
    bg, fg = BLOOM_COLOURS.get(key, ("#F3F4F6", "#374151"))
    return (
        f'<span style="display:inline-block;padding:2px 10px;border-radius:20px;'
        f'font-size:0.72rem;font-weight:700;background:{bg};color:{fg};'
        f'margin-left:6px;vertical-align:middle;">🧠 {level}</span>'
    )

def render_exam_html(text: str, is_key=False) -> str:
    accent = AMBER if is_key else TEAL
    lines  = text.split("\n")
    out    = []
    for line in lines:
        s = line.strip()
        if not s or s in ("---EXAM START---", "---ANSWER KEY---"):
            out.append("<br>")
            continue

        def replace_bloom(m):
            return bloom_badge_html(m.group(1))
        s_rendered = BLOOM_RE.sub(replace_bloom, s)

        if s.startswith("## "):
            t = s[3:].strip()
            out.append(
                f'<h2 style="color:{accent};border-bottom:2px solid {TEAL_LIGHT};'
                f'padding-bottom:6px;margin-top:1.4rem;">{BLOOM_RE.sub(replace_bloom, t)}</h2>'
            )
        elif s.startswith("### "):
            t = s[4:].strip()
            out.append(
                f'<h3 style="color:{SLATE};background:{TEAL_LIGHT};padding:8px 14px;'
                f'border-radius:6px;border-left:4px solid {accent};margin-top:1.2rem;">'
                f'{BLOOM_RE.sub(replace_bloom, t)}</h3>'
            )
        elif s.startswith("#### "):
            t = s[5:].strip()
            out.append(f'<h4 style="color:{AMBER};margin-top:1rem;">{BLOOM_RE.sub(replace_bloom, t)}</h4>')
        elif re.match(r'^[A-D]\)', s):
            out.append(
                f'<div style="padding:3px 0 3px 2rem;color:{SLATE_LIGHT};font-size:0.95rem;">{s}</div>'
            )
        elif re.match(r'^Q\d+\.', s):
            out.append(
                f'<p style="font-weight:700;color:{SLATE};margin:1.1rem 0 0.3rem 0;'
                f'font-size:1rem;">{s_rendered}</p>'
            )
        elif s.startswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):
                continue
            row_html = "".join(
                f'<td style="padding:6px 10px;border:1px solid {BORDER};color:{SLATE};">{c}</td>'
                for c in cells
            )
            bg = TEAL_LIGHT if not is_key else AMBER_LIGHT
            out.append(f'<table style="width:100%;border-collapse:collapse;margin:6px 0;">'
                       f'<tr style="background:{bg};">{row_html}</tr></table>')
        elif s.startswith("---"):
            out.append(f'<hr style="border:1px solid {BORDER};margin:1rem 0;">')
        else:
            out.append(f'<p style="color:{SLATE};margin:0.3rem 0;">{s_rendered}</p>')
    return "\n".join(out)


# ─────────────────────────────────────────────
# PDF EXPORT
# ─────────────────────────────────────────────
def _rl_styles():
    S = {}
    S["title"]    = ParagraphStyle("T",  fontName="Helvetica-Bold",   fontSize=22, textColor=RL_TEAL,       alignment=TA_CENTER, spaceAfter=6)
    S["subtitle"] = ParagraphStyle("ST", fontName="Helvetica",        fontSize=10, textColor=RL_SLATE_LIGHT, alignment=TA_CENTER, spaceAfter=18)
    S["h2"]       = ParagraphStyle("H2", fontName="Helvetica-Bold",   fontSize=14, textColor=RL_TEAL,       spaceBefore=14, spaceAfter=5)
    S["h2key"]    = ParagraphStyle("H2K",fontName="Helvetica-Bold",   fontSize=14, textColor=RL_AMBER,      spaceBefore=14, spaceAfter=5)
    S["h3"]       = ParagraphStyle("H3", fontName="Helvetica-Bold",   fontSize=11, textColor=RL_SLATE,      spaceBefore=10, spaceAfter=4, backColor=RL_TEAL_LIGHT, leftIndent=8, borderPad=6)
    S["h4"]       = ParagraphStyle("H4", fontName="Helvetica-Bold",   fontSize=10, textColor=RL_AMBER,      spaceBefore=8, spaceAfter=3)
    S["bloom"]    = ParagraphStyle("BL", fontName="Helvetica-BoldOblique", fontSize=8, textColor=RL_TEAL, spaceAfter=1)
    S["question"] = ParagraphStyle("Q",  fontName="Helvetica-Bold",   fontSize=10, textColor=RL_SLATE,      spaceBefore=8, spaceAfter=3, leading=14)
    S["option"]   = ParagraphStyle("OP", fontName="Helvetica",        fontSize=10, textColor=RL_SLATE_LIGHT, leftIndent=22, spaceAfter=2, leading=13)
    S["body"]     = ParagraphStyle("BD", fontName="Helvetica",        fontSize=10, textColor=RL_SLATE,      spaceAfter=4, leading=15, alignment=TA_JUSTIFY)
    S["keybody"]  = ParagraphStyle("KB", fontName="Helvetica",        fontSize=10, textColor=RL_SLATE,      spaceAfter=4, leading=14, backColor=RL_AMBER_LIGHT, leftIndent=8, borderPad=4)
    return S

def build_pdf(full_text, title, grade, difficulty, duration, total_marks) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2.2*cm, bottomMargin=2*cm)
    S     = _rl_styles()
    story = []

    # Header
    story.append(Paragraph("ExamCraft AI", S["title"]))
    story.append(Paragraph(title, S["subtitle"]))

    # Meta table
    meta = [["Level", grade.split("(")[0].strip(), "Difficulty", difficulty],
            ["Duration", f"{duration} mins", "Total Marks", str(total_marks)]]
    mt = Table(meta, colWidths=[2.5*cm, 6.5*cm, 2.5*cm, 3.5*cm])
    mt.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (0,-1), RL_TEAL),
        ("BACKGROUND",   (2,0), (2,-1), RL_TEAL),
        ("TEXTCOLOR",    (0,0), (0,-1), RL_WHITE),
        ("TEXTCOLOR",    (2,0), (2,-1), RL_WHITE),
        ("BACKGROUND",   (1,0), (1,-1), RL_TEAL_LIGHT),
        ("BACKGROUND",   (3,0), (3,-1), RL_TEAL_LIGHT),
        ("TEXTCOLOR",    (1,0), (3,-1), RL_SLATE),
        ("FONTNAME",     (0,0), (-1,-1), "Helvetica"),
        ("FONTNAME",     (0,0), (0,-1), "Helvetica-Bold"),
        ("FONTNAME",     (2,0), (2,-1), "Helvetica-Bold"),
        ("FONTSIZE",     (0,0), (-1,-1), 9),
        ("ALIGN",        (0,0), (-1,-1), "CENTER"),
        ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
        ("BOX",          (0,0), (-1,-1), 0.5, RL_BORDER),
        ("INNERGRID",    (0,0), (-1,-1), 0.3, RL_BORDER),
        ("TOPPADDING",   (0,0), (-1,-1), 6),
        ("BOTTOMPADDING",(0,0), (-1,-1), 6),
    ]))
    story.append(mt)
    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=2, color=RL_TEAL))
    story.append(Spacer(1, 8))

    exam_body, answer_key = split_exam_and_key(full_text)

    def render(text, is_key=False):
        for line in text.split("\n"):
            s = line.strip()
            if not s or s in ("---EXAM START---", "---ANSWER KEY---"):
                story.append(Spacer(1, 4))
                continue

            bloom_m = BLOOM_RE.search(s)
            bloom_l = bloom_m.group(1) if bloom_m else None
            s_clean = BLOOM_RE.sub("", s).strip()

            if s.startswith("## "):
                style = S["h2key"] if is_key else S["h2"]
                story.append(Paragraph(s[3:].strip(), style))
                story.append(HRFlowable(width="100%", thickness=1,
                                        color=RL_AMBER if is_key else RL_TEAL))
            elif s.startswith("### "):
                story.append(Paragraph(s[4:].strip(), S["h3"]))
            elif s.startswith("#### "):
                story.append(Paragraph(s[5:].strip(), S["h4"]))
            elif re.match(r'^[A-D]\)', s_clean):
                story.append(Paragraph(s_clean, S["option"]))
            elif re.match(r'^Q\d+\.', s_clean):
                if bloom_l:
                    story.append(Paragraph(f"Bloom's Taxonomy: {bloom_l}", S["bloom"]))
                story.append(Paragraph(s_clean, S["question"]))
            elif s.startswith("|"):
                cells = [c.strip() for c in s.strip("|").split("|")]
                if all(set(c) <= set("-: ") for c in cells):
                    continue
                tbl = Table([cells])
                bg  = RL_AMBER_LIGHT if is_key else RL_TEAL_LIGHT
                tbl.setStyle(TableStyle([
                    ("BACKGROUND",   (0,0), (-1,-1), bg),
                    ("TEXTCOLOR",    (0,0), (-1,-1), RL_SLATE),
                    ("FONTNAME",     (0,0), (-1,-1), "Helvetica"),
                    ("FONTSIZE",     (0,0), (-1,-1), 8),
                    ("BOX",          (0,0), (-1,-1), 0.4, RL_BORDER),
                    ("INNERGRID",    (0,0), (-1,-1), 0.2, RL_BORDER),
                    ("TOPPADDING",   (0,0), (-1,-1), 4),
                    ("BOTTOMPADDING",(0,0), (-1,-1), 4),
                    ("ALIGN",        (0,0), (-1,-1), "LEFT"),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 2))
            elif s.startswith("---"):
                story.append(Spacer(1, 4))
            else:
                style = S["keybody"] if is_key else S["body"]
                story.append(Paragraph(s_clean or s, style))

    render(exam_body, is_key=False)
    if answer_key:
        story.append(PageBreak())
        story.append(Paragraph("ANSWER KEY & GRADING RUBRIC", S["h2key"]))
        story.append(HRFlowable(width="100%", thickness=2, color=RL_AMBER))
        story.append(Spacer(1, 8))
        render(answer_key, is_key=True)

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# DOCX EXPORT
# ─────────────────────────────────────────────
def build_docx(full_text, title) -> bytes:
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r  = tp.add_run("ExamCraft AI — " + title)
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = D_TEAL
    doc.add_paragraph()

    exam_body, answer_key = split_exam_and_key(full_text)

    def add_sec(text, is_key=False):
        accent = D_AMBER if is_key else D_TEAL
        for line in text.split("\n"):
            s = line.strip()
            if not s or s in ("---EXAM START---", "---ANSWER KEY---"):
                doc.add_paragraph(); continue

            bloom_m = BLOOM_RE.search(s)
            bloom_l = bloom_m.group(1) if bloom_m else None
            s_clean = BLOOM_RE.sub("", s).strip()

            if s.startswith("## "):
                h = doc.add_heading(s[3:].strip(), level=1)
                for r in h.runs: r.font.color.rgb = accent
            elif s.startswith("### "):
                h = doc.add_heading(s[4:].strip(), level=2)
                for r in h.runs: r.font.color.rgb = accent
            elif s.startswith("#### "):
                p = doc.add_paragraph()
                r = p.add_run(s[5:].strip())
                r.bold = True; r.font.size = Pt(11); r.font.color.rgb = D_AMBER
            elif re.match(r'^[A-D]\)', s_clean):
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(s_clean); r.font.color.rgb = D_GRAY
            elif re.match(r'^Q\d+\.', s_clean):
                if bloom_l:
                    bp = doc.add_paragraph()
                    br = bp.add_run(f"  Bloom's Taxonomy: {bloom_l}")
                    br.italic = True; br.font.size = Pt(8); br.font.color.rgb = D_TEAL
                p = doc.add_paragraph()
                r = p.add_run(s_clean)
                r.bold = True; r.font.size = Pt(11); r.font.color.rgb = D_SLATE
            elif s.startswith("|"):
                cells = [c.strip() for c in s.strip("|").split("|")]
                if all(set(c) <= set("-: ") for c in cells): continue
                p = doc.add_paragraph("  |  ".join(cells))
                p.paragraph_format.left_indent = Inches(0.15)
            elif s.startswith("---"):
                doc.add_paragraph()
            else:
                p = doc.add_paragraph(s_clean or s)
                if is_key: p.paragraph_format.left_indent = Inches(0.15)

    add_sec(exam_body, is_key=False)
    if answer_key:
        doc.add_page_break()
        h = doc.add_heading("ANSWER KEY & GRADING RUBRIC", level=1)
        for r in h.runs: r.font.color.rgb = D_AMBER
        add_sec(answer_key, is_key=True)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════
# ══  UI  ═══════════════════════════════════════
# ═══════════════════════════════════════════════

# ── Sidebar ──
with st.sidebar:
    st.markdown(f"""
    <div style="padding:1.5rem 0 1rem 0;border-bottom:1px solid rgba(26,122,110,0.3);margin-bottom:1.2rem;">
        <div style="font-size:0.65rem;font-weight:800;letter-spacing:3px;text-transform:uppercase;color:{TEAL_MID};margin-bottom:0.3rem;">ExamCraft AI</div>
        <div style="font-size:1.3rem;font-weight:800;color:{WHITE};letter-spacing:-0.5px;">Configuration</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown(f'<div style="font-size:0.68rem;font-weight:700;letter-spacing:2px;text-transform:uppercase;color:{TEAL_MID};margin-bottom:0.5rem;">🎓 Academic Level</div>', unsafe_allow_html=True)
    grade = st.selectbox("", [
        "Grade 6–8 (Middle School)", "Grade 9–10 (Secondary)",
        "Grade 11–12 (Pre-University)", "Semester 1 (University)",
        "Semester 2 (University)", "Semester 3–4 (University)", "Postgraduate Level",
    ], label_visibility="collapsed")

    st.markdown(f'<div style="font-size:0.68rem;font-weight:700;letter-spacing:2px;text-transform:uppercase;color:{TEAL_MID};margin:1rem 0 0.5rem 0;">📊 Difficulty</div>', unsafe_allow_html=True)
    difficulty = st.selectbox("", ["Beginner", "Intermediate", "Advanced", "Expert"], index=1, label_visibility="collapsed")

    st.markdown(f'<div style="font-size:0.68rem;font-weight:700;letter-spacing:2px;text-transform:uppercase;color:{TEAL_MID};margin:1rem 0 0.5rem 0;">⏱️ Duration & Marks</div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1: duration    = st.number_input("Duration (min)", min_value=15, max_value=300, value=60,  step=15)
    with c2: total_marks = st.number_input("Total Marks",    min_value=10, max_value=200, value=100, step=5)

    st.markdown(f"""
    <div style="border-top:1px solid rgba(26,122,110,0.3);margin:1.4rem 0 1rem 0;padding-top:1.2rem;">
        <div style="font-size:0.68rem;font-weight:700;letter-spacing:2px;text-transform:uppercase;color:{TEAL_MID};margin-bottom:0.8rem;">📝 Question Counts</div>
    </div>
    """, unsafe_allow_html=True)

    num_mcq   = st.number_input("🔵 MCQs",             min_value=0, max_value=30, value=5,  step=1)
    num_short = st.number_input("🟠 Short Questions",  min_value=0, max_value=20, value=3,  step=1)
    num_long  = st.number_input("🟣 Long / Essay Qs",  min_value=0, max_value=10, value=1,  step=1)

    st.markdown(f"""
    <div style="border-top:1px solid rgba(26,122,110,0.3);margin:1.4rem 0 1rem 0;padding-top:1.2rem;">
        <div style="font-size:0.68rem;font-weight:700;letter-spacing:2px;text-transform:uppercase;color:{TEAL_MID};margin-bottom:0.8rem;">🧠 Bloom's Taxonomy</div>
    </div>
    """, unsafe_allow_html=True)
    badges_html = "".join(
        f'<span style="background:{bg};color:{fg};padding:3px 10px;border-radius:20px;'
        f'font-size:0.7rem;font-weight:700;display:inline-block;margin:3px 2px;">'
        f'{lvl.capitalize()}</span>'
        for lvl, (bg, fg) in BLOOM_COLOURS.items()
    )
    st.markdown(
        f'<div style="background:rgba(255,255,255,0.06);border:1px solid rgba(26,122,110,0.25);border-radius:10px;'
        f'padding:10px;line-height:2.2;">{badges_html}</div>',
        unsafe_allow_html=True,
    )
    st.markdown("<br>", unsafe_allow_html=True)
    st.caption("ExamCraft AI · Groq LLaMA 3.1 · v2.0")


# ── Hero ──
st.markdown(f"""
<div class="hero-banner">
  <div class="hero-eyebrow">AI-Powered Assessment Tool</div>
  <div class="hero-title">Exam<span>Craft</span> AI</div>
  <div class="hero-sub">Generate structured, Bloom's-tagged exam papers in seconds — export to PDF, DOCX & TXT with one click.</div>
</div>
""", unsafe_allow_html=True)


# ── Step 1 ──
st.markdown('<div class="section-card">', unsafe_allow_html=True)
st.markdown('<div class="section-label">Step 01</div>', unsafe_allow_html=True)
st.markdown('<div class="section-title">📂 Provide Exam Content</div>', unsafe_allow_html=True)

# Question count summary
total_q = num_mcq + num_short + num_long
st.markdown(f"""
<div class="q-counter-grid">
  <div class="q-counter-card mcq">
    <div class="q-counter-label">MCQs</div>
    <div style="font-size:2rem;font-weight:900;color:{TEAL};line-height:1;">{num_mcq}</div>
    <div style="font-size:0.72rem;color:{SLATE_LIGHT};margin-top:2px;">questions</div>
  </div>
  <div class="q-counter-card short">
    <div class="q-counter-label">Short Qs</div>
    <div style="font-size:2rem;font-weight:900;color:{AMBER};line-height:1;">{num_short}</div>
    <div style="font-size:0.72rem;color:{SLATE_LIGHT};margin-top:2px;">questions</div>
  </div>
  <div class="q-counter-card long">
    <div class="q-counter-label">Long / Essay</div>
    <div style="font-size:2rem;font-weight:900;color:#7E22CE;line-height:1;">{num_long}</div>
    <div style="font-size:0.72rem;color:{SLATE_LIGHT};margin-top:2px;">questions</div>
  </div>
</div>
<div style="text-align:right;font-size:0.78rem;color:{SLATE_LIGHT};margin-top:-0.5rem;margin-bottom:1rem;font-weight:600;">
  Total: <span style="color:{TEAL};font-weight:800;">{total_q} questions</span> · Adjust counts in the sidebar →
</div>
""", unsafe_allow_html=True)

source_type   = st.radio("How would you like to provide the topic?",
                         ["Enter Topic Manually", "Upload Source Material"], horizontal=True)
topic_context = ""

if source_type == "Enter Topic Manually":
    topic_context = st.text_area(
        "Subject & Topics",
        placeholder="e.g., Subject: Physics\nTopics: Thermodynamics, Laws of Entropy, Heat Transfer",
        height=140,
    )
else:
    st.markdown(f"""
    <div style="background:{TEAL_LIGHT};border:1.5px dashed {TEAL};border-radius:12px;
    padding:0.6rem 1rem;margin-bottom:0.6rem;">
      <span style="color:{TEAL};font-weight:700;font-size:0.88rem;">📎 Accepted formats: PDF, TXT · Max 200MB per file</span>
    </div>
    """, unsafe_allow_html=True)
    uploaded_file = st.file_uploader("Upload syllabus, notes, or chapter", type=["pdf","txt"], label_visibility="collapsed")
    if uploaded_file:
        with st.spinner("Extracting content…"):
            topic_context = extract_file_content(uploaded_file)
        if topic_context:
            st.success(f"✅ Extracted **{len(topic_context.split())} words** from `{uploaded_file.name}`")
            with st.expander("🔍 Preview extracted content"):
                st.markdown(f"""
                <div style="background:#F8FAFA;border:1px solid {BORDER};border-radius:8px;
                padding:1rem 1.2rem;color:{SLATE};font-size:0.85rem;line-height:1.7;
                white-space:pre-wrap;font-family:monospace;">{topic_context[:1500] + ("…" if len(topic_context) > 1500 else "")}</div>
                """, unsafe_allow_html=True)
        else:
            st.warning("Could not extract text from this file.")
st.markdown('</div>', unsafe_allow_html=True)


# ── Step 2 ──
st.markdown('<div class="section-card">', unsafe_allow_html=True)
st.markdown('<div class="section-label">Step 02</div>', unsafe_allow_html=True)
st.markdown('<div class="section-title">⚡ Generate Your Exam</div>', unsafe_allow_html=True)
st.markdown(f'<p style="color:{SLATE_LIGHT};font-size:0.9rem;margin-bottom:1rem;">AI will craft {num_mcq} MCQs, {num_short} short questions, and {num_long} long/essay questions — all Bloom\'s tagged.</p>', unsafe_allow_html=True)
gen_col, _ = st.columns([1, 3])
with gen_col:
    generate_btn = st.button("⚡ Generate Exam Paper", use_container_width=True)
st.markdown('</div>', unsafe_allow_html=True)


# ── Generation ──
if generate_btn:
    if not topic_context.strip():
        st.warning("⚠️ Please enter a topic or upload a source file first.")
        st.stop()
    api_key = load_api_key()
    prompt  = build_prompt(topic_context, grade, difficulty, duration, total_marks, source_type, num_mcq, num_short, num_long)
    with st.spinner("🧠 Crafting your exam paper — this may take 20–40 seconds…"):
        try:
            full_output = generate_exam(prompt, api_key)
            st.session_state.update({
                "full_output": full_output, "exam_grade": grade,
                "exam_diff": difficulty, "exam_dur": duration, "exam_marks": total_marks,
                "exam_mcq": num_mcq, "exam_short": num_short, "exam_long": num_long,
            })
        except Exception as e:
            st.error(f"❌ Generation failed: {e}")
            st.stop()


# ── Results ──
if "full_output" in st.session_state:
    full_output = st.session_state["full_output"]
    grade_label = st.session_state.get("exam_grade",  grade)
    diff_label  = st.session_state.get("exam_diff",   difficulty)
    dur_val     = st.session_state.get("exam_dur",    duration)
    marks_val   = st.session_state.get("exam_marks",  total_marks)

    exam_body, answer_key = split_exam_and_key(full_output)

    st.success("✅ Exam generated successfully!")

    mc1, mc2, mc3, mc4, mc5, mc6 = st.columns(6)
    mc1.metric("Level",        grade_label.split("(")[0].strip())
    mc2.metric("Difficulty",   diff_label)
    mc3.metric("Duration",     f"{dur_val} mins")
    mc4.metric("Total Marks",  str(marks_val))
    mc5.metric("Questions",    str(st.session_state.get("exam_mcq",0) + st.session_state.get("exam_short",0) + st.session_state.get("exam_long",0)))
    mc6.metric("Sections",     "3")

    st.markdown("---")
    st.markdown("## 📄 Generated Exam Paper")
    st.markdown(
        f'<div class="exam-output">{render_exam_html(exam_body, is_key=False)}</div>',
        unsafe_allow_html=True,
    )

    if answer_key:
        st.markdown("---")
        with st.expander("🔑 View Answer Key & Grading Rubric (Teacher Only)", expanded=False):
            st.info("📌 Keep this section hidden when distributing to students.")
            st.markdown(
                f'<div class="exam-output" style="border-left-color:{AMBER};">'
                f'{render_exam_html(answer_key, is_key=True)}</div>',
                unsafe_allow_html=True,
            )

    # ── Step 3: Export ──
    st.markdown("---")
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">Step 03</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">💾 Export Your Exam</div>', unsafe_allow_html=True)

    exam_title  = f"Exam — {grade_label} — {diff_label}"
    txt_content = f"{exam_title}\n{'='*60}\n\n{full_output}"

    with st.spinner("Preparing download files…"):
        docx_bytes = build_docx(full_output, exam_title)
        pdf_bytes  = build_pdf(full_output, exam_title, grade_label, diff_label, dur_val, marks_val)

    dl1, dl2, dl3 = st.columns(3)
    with dl1:
        st.download_button("📄 Download .txt", data=txt_content.encode("utf-8"),
                           file_name="exam_paper.txt", mime="text/plain",
                           use_container_width=True)
    with dl2:
        st.download_button("📝 Download .docx", data=docx_bytes,
                           file_name="exam_paper.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
    with dl3:
        st.download_button("🖨️ Download .pdf", data=pdf_bytes,
                           file_name="exam_paper.pdf", mime="application/pdf",
                           use_container_width=True)

    st.caption("All formats include the exam + answer key. PDF and DOCX use teal/amber theme with Bloom's labels.")
    st.markdown('</div>', unsafe_allow_html=True)
